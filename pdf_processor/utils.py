"""
Core PDF processing utilities converted from Flask app
"""
import os
import sys
import logging
from pathlib import Path
import json
import base64
import time
import gc
from io import BytesIO
from concurrent.futures import ThreadPoolExecutor
from tenacity import retry, stop_after_attempt, wait_exponential

# Load environment variables from .env file
try:
    from dotenv import load_dotenv, find_dotenv
    load_dotenv(find_dotenv(), override=False)
except ImportError:
    logging.warning("python-dotenv not available. Environment variables must be set manually.")

# --- Windows DLL bootstrap (must run BEFORE importing pdfixsdk) ---
if os.name == "nt":
    import importlib.util

    try:
        spec = importlib.util.find_spec("pdfixsdk")
        dll_dir = None
        if spec and spec.submodule_search_locations:
            pkg_dir = Path(list(spec.submodule_search_locations)[0])
            for cand in [
                pkg_dir / "bin" / "x86_64" / "windows",
                pkg_dir / "bin" / "x86_64",
                pkg_dir / "bin",
            ]:
                if cand.exists():
                    dll_dir = cand
                    break
        if dll_dir:
            os.add_dll_directory(str(dll_dir))
            os.environ["PATH"] = str(dll_dir) + ";" + os.environ.get("PATH", "")
            logging.info(f"PDFix DLL dir added early: {dll_dir}")
    except Exception as e:
        logging.info(f"PDFix DLL early-path setup failed: {e}")

import fitz  # PyMuPDF
from PIL import Image
try:
    from pdfixsdk import *
    PDFIX_AVAILABLE = True
    logging.info("PDFix SDK loaded successfully")
except ImportError as e:
    PDFIX_AVAILABLE = False
    logging.warning(f"PDFix SDK not available ({e}). PyMuPDF fallback will be used.")
except Exception as e:
    PDFIX_AVAILABLE = False
    logging.warning(f"PDFix SDK initialization failed ({e}). PyMuPDF fallback will be used.")
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
import psutil
from django.conf import settings

# Soft memory guard
PROC = psutil.Process(os.getpid())
SOFT_MEM_LIMIT_MB = float(os.getenv("SOFT_MEM_LIMIT_MB", "1700"))
PDFIX_RENDER_RESOLUTION = float(os.getenv("PDFIX_RENDER_RESOLUTION", "1.0"))
MAX_ALT_TEXT_WORKERS = int(os.getenv("MAX_ALT_TEXT_WORKERS", 4))
if os.getenv("RENDER") == "true":
    MAX_ALT_TEXT_WORKERS = 2

from threading import Lock

progress_cache = {}
progress_lock = Lock()

# Get sessions directory from Django settings or use default
def get_sessions_dir():
    return Path(getattr(settings, 'PDF_SESSIONS_DIR', 'sessions'))


def rss_mb() -> float:
    return PROC.memory_info().rss / (1024 * 1024)


def validate_session(sid: str) -> bool:
    """Check if a session directory exists for the given sid."""
    session_dir = get_sessions_dir() / sid
    try:
        return session_dir.exists() and session_dir.is_dir()
    except Exception as e:
        logging.info(f"Session validation failed for sid {sid}: {e}")
        return False


def sess_paths(sid: str):
    """Get all paths for a session."""
    base = get_sessions_dir() / sid
    return {
        "base": base,
        "uploads": base / "uploads",
        "extracted": base / "extracted",
        "output": base / "output",
        "img_progress": base / "progress_image.txt",
        "alt_progress": base / "progress_alt.txt",
        "page_map": base / "page_numbers.json",
        "img_status": base / "status_image.txt",
    }


def ensure_dirs(p):
    """Ensure all session directories exist."""
    p["base"].mkdir(parents=True, exist_ok=True)
    p["uploads"].mkdir(parents=True, exist_ok=True)
    p["extracted"].mkdir(parents=True, exist_ok=True)
    p["output"].mkdir(parents=True, exist_ok=True)


def write_json(path: Path, obj):
    """Write JSON to file atomically."""
    path.write_text(json.dumps(obj, ensure_ascii=False), encoding="utf-8")


def read_json(path: Path, default):
    """Read JSON from file with default fallback."""
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def write_progress(path: Path, value: int):
    """Write progress value to file with caching."""
    with progress_lock:
        progress_cache[path] = value
        if value % 5 == 0 or value >= 100 or value < 0:
            try:
                tmp = path.with_suffix(path.suffix + ".tmp")
                tmp.write_text(str(int(value)), encoding="utf-8")
                os.replace(tmp, path)
                logging.info(f"Wrote progress to {path}: {value}%")
            except Exception as e:
                logging.info(f"Failed to write progress to {path}: {e}")


def read_progress(path: Path) -> int:
    """Read progress value from file."""
    try:
        return int(path.read_text(encoding="utf-8").strip())
    except Exception:
        try:
            time.sleep(0.05)
            return int(path.read_text(encoding="utf-8").strip())
        except Exception:
            return 0


def write_status(path: Path, text: str):
    """Write status text to file."""
    try:
        tmp = path.with_suffix(path.suffix + ".tmp")
        tmp.write_text(text.strip(), encoding="utf-8")
        os.replace(tmp, path)
    except Exception as e:
        logging.info(f"❌ Failed to write status to {path}: {e}")


def read_status(path: Path) -> str:
    """Read status text from file."""
    try:
        return path.read_text(encoding="utf-8").strip()
    except Exception:
        return ""


def write_alt_progress_detail(path: Path, percent: int, done: int, total: int):
    """Write detailed alt text progress."""
    with progress_lock:
        key = str(path)
        if key not in progress_cache or progress_cache[key]["percent"] != percent:
            progress_cache[key] = {"percent": percent, "done": done, "total": total}
            try:
                tmp = path.with_suffix(path.suffix + ".tmp")
                tmp.write_text(
                    f"{int(percent)}|{int(done)}|{int(total)}",
                    encoding="utf-8",
                )
                os.replace(tmp, path)
                logging.info(
                    f"Wrote alt progress to {path}: {percent}% ({done}/{total})"
                )
            except Exception as e:
                logging.info(f"Failed to write alt progress to {path}: {e}")


def read_alt_progress_detail(path: Path):
    """Read detailed alt text progress."""
    with progress_lock:
        key = str(path)
        if key in progress_cache:
            return (
                progress_cache[key]["percent"],
                progress_cache[key]["done"],
                progress_cache[key]["total"],
            )
        try:
            raw = path.read_text(encoding="utf-8").strip()
            parts = raw.split("|")
            if len(parts) >= 3:
                p, d, t = parts[:3]
                return int(p or 0), int(d or 0), int(t or 0)
            return int(raw or 0), 0, 0
        except Exception:
            return 0, 0, 0


_openai_client = None


def get_openai_client():
    """Return a singleton OpenAI client."""
    global _openai_client
    if _openai_client is None:
        key = os.getenv("OPENAI_API_KEY") or os.getenv("OPEN_API_KEY") or ""
        key = key.strip().strip('"').strip("'")
        if not key or not key.startswith("sk-"):
            logging.error(
                "No valid OpenAI API key found (OPENAI_API_KEY / OPEN_API_KEY)."
            )
        _openai_client = OpenAI(api_key=key)
    return _openai_client


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=4, max=10))
def call_alttext_ai(image_path, lang="en"):
    """Generate alt text for an image using OpenAI."""
    try:
        img = Image.open(image_path)
        max_dim = 1024
        if max(img.size) > max_dim:
            img.thumbnail((max_dim, max_dim), Image.LANCZOS)

        buf = BytesIO()
        img.save(buf, format="PNG")
        img_data = buf.getvalue()
        buf.close()
        img_b64 = base64.b64encode(img_data).decode()

        if lang == "en":
            prompt = "Generate concise alt text for this image. Prioritize clarity and brevity unless the image contains complex information like graphs or diagrams. Omit unnecessary visual details."
        elif lang == "es":
            prompt = "Genera un texto alternativo conciso para esta imagen. Prioriza la claridad y la brevedad a menos que la imagen contenga información compleja como gráficos o diagramas. Omite detalles visuales innecesarios."
        else:
            prompt = f"Generate alt text in {lang} using the same rules: be concise unless the image is a chart or diagram."

        client = get_openai_client()

        # Try gpt-4o-mini first (cheaper), fallback to gpt-4o if needed
        try:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert in accessibility, generating alt text for images.",
                    },
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": "data:image/png;base64," + img_b64
                                },
                            },
                        ],
                    },
                ],
                max_tokens=150,
            )
            text = response.choices[0].message.content
            return (text[0] if isinstance(text, list) and text else text).strip()
        except Exception as e:
            logging.info(
                f"OpenAI call failed (gpt-4o-mini). Trying gpt-4o: {e}"
            )
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    messages=[
                        {
                            "role": "system",
                            "content": "You are an expert in accessibility, generating alt text for images.",
                        },
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": prompt},
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": "data:image/png;base64," + img_b64
                                    },
                                },
                            ],
                        },
                    ],
                    max_tokens=150,
                )
                text = response.choices[0].message.content
                return (text[0] if isinstance(text, list) and text else text).strip()
            except Exception as e2:
                logging.error(f"Both models failed. Last error: {e2}")
                raise e
    except Exception as e:
        logging.info(f"Failed to generate alt text via OpenAI: {e}")
        return f"[Error: {str(e)}]"


def extract_images_from_pdfix_per_session(
    pdf_path, output_folder, progress_file: Path
):
    """Extract images from PDF using PDFix SDK."""
    if not PDFIX_AVAILABLE:
        raise ImportError("PDFix SDK not available. Use PyMuPDF fallback instead.")
    pdfix = GetPdfix()
    if not pdfix:
        raise Exception("PDFix initialization failed!")

    doc = pdfix.OpenDoc(pdf_path, "")
    if not doc:
        raise Exception("Unable to open PDF document.")

    os.makedirs(output_folder, exist_ok=True)
    image_paths = []
    page_numbers = {}

    num_pages = doc.GetNumPages()
    scale = float(PDFIX_RENDER_RESOLUTION)

    def save_image(page, element, page_num, page_view):
        nonlocal image_paths, page_numbers
        if element.GetType() == kPdeImage:
            image = PdeImage(element.obj)
            bbox = image.GetBBox()

            dev = page_view.RectToDevice(bbox)
            dev_w = max(1, dev.right - dev.left)
            dev_h = max(1, dev.bottom - dev.top)
            max_pixels = int(os.getenv("MAX_BBOX_PIXELS", "50000000"))
            if dev_w * dev_h > max_pixels:
                raise MemoryError(f"bbox too large: {dev_w}x{dev_h} ({dev_w*dev_h} px)")

            ps_image = pdfix.CreateImage(dev_w, dev_h, kImageDIBFormatArgb)

            m = page_view.GetDeviceMatrix()
            try:
                m.e -= dev.left
                m.f -= dev.top
            except Exception:
                try:
                    m = PdfMatrix(m.a, m.b, m.c, m.d, m.e - dev.left, m.f - dev.top)
                except Exception:
                    pass

            params = PdfPageRenderParams()
            params.clip_box = bbox
            params.image = ps_image
            params.matrix = m
            page.DrawContent(params)

            dev.left, dev.top, dev.right, dev.bottom = 0, 0, dev_w, dev_h

            img_path = os.path.join(
                output_folder, f"Extracted_Image_{len(image_paths)+1}.png"
            )
            ps_image.SaveRect(img_path, PdfImageParams(), dev)
            ps_image.Destroy()

            fn = os.path.basename(img_path)
            image_paths.append({"path": img_path, "filename": fn, "page_num": page_num})
            page_numbers[fn] = page_num
        else:
            for i in range(element.GetNumChildren()):
                child = element.GetChild(i)
                if child:
                    save_image(page, child, page_num, page_view)

    for i in range(num_pages):
        if rss_mb() > SOFT_MEM_LIMIT_MB:
            raise MemoryError(
                f"Soft memory cap exceeded before page {i+1}: {rss_mb():.0f} MB"
            )

        page = doc.AcquirePage(i)
        page_map = page.AcquirePageMap()
        if not page_map or not page_map.CreateElements():
            try:
                page.Release()
            finally:
                raise Exception("Failed to acquire PageMap.")
        container = page_map.GetElement()
        if not container:
            try:
                page_map.Release()
            except Exception:
                try:
                    page_map.Destroy()
                except Exception:
                    pass
            page.Release()
            raise Exception("Failed to get page elements.")

        if i == 0:
            try:
                status_path = Path(progress_file).with_name("status_image.txt")
                write_status(status_path, "Extracting images…")
            except Exception:
                pass

        if rss_mb() > (SOFT_MEM_LIMIT_MB - 500) and scale > 0.45:
            new_scale = max(0.4, round(scale - 0.1, 2))
            logging.info(
                f"[mem] RSS={rss_mb():.0f}MB near limit; lowering scale {scale}→{new_scale}"
            )
            scale = new_scale

        page_view = page.AcquirePageView(scale, kRotate0)
        save_image(page, container, i + 1, page_view)
        write_progress(progress_file, int(((i + 1) / num_pages) * 100))

        try:
            page_view.Release()
        except Exception:
            try:
                page_view.Destroy()
            except Exception:
                pass
        try:
            page_map.Release()
        except Exception:
            try:
                page_map.Destroy()
            except Exception:
                pass
        page.Release()

        if (i % 8) == 7:
            gc.collect()
            logging.info(
                f"[mem] after GC: RSS={rss_mb():.0f}MB at page {i+1}/{num_pages}"
            )

    doc.Close()
    del doc, pdfix
    gc.collect()
    return image_paths, page_numbers


def extract_images_with_pymupdf(
    pdf_path, output_folder, progress_file: Path, sid: str = None
):
    """Extract images using PyMuPDF as fallback."""
    os.makedirs(output_folder, exist_ok=True)
    image_paths, page_numbers = [], {}
    doc = fitz.open(pdf_path)
    n = doc.page_count
    write_progress(progress_file, 0)
    for i in range(n):
        logging.info(f"[sid={sid}] (fitz) page {i+1}/{n}")
        page = doc.load_page(i)
        for _, img in enumerate(page.get_images(full=True), start=1):
            xref = img[0]
            try:
                pix = fitz.Pixmap(doc, xref)
                if pix.alpha:
                    pix = fitz.Pixmap(pix, 0)
                out = os.path.join(
                    output_folder, f"Extracted_Image_{len(image_paths)+1}.png"
                )
                pix.save(out)
                fn = os.path.basename(out)
                image_paths.append({"path": out, "filename": fn, "page_num": i + 1})
                page_numbers[fn] = i + 1
            except Exception:
                continue
        write_progress(progress_file, int(((i + 1) / max(1, n)) * 100))
    doc.close()
    return image_paths, page_numbers


def _write_copy_panel_html(output_path: str, entries, title="Alt Text Copy Panel"):
    """Generate HTML copy panel."""
    import html

    rows = []
    for i, e in enumerate(entries, 1):
        alt = html.escape(e.get("alt", ""))
        page = e.get("page")
        thumb = e.get("thumb_rel")
        page_html = f'<span class="page">Page {page}</span>' if page else ""
        thumb_html = f'<img src="{thumb}" alt="" class="thumb"/>' if thumb else ""
        rows.append(
            f"""
        <div class="item">
          <div class="meta">
            {thumb_html}
            {page_html}
          </div>
          <pre id="alt{i}" class="alt">{alt}</pre>
          <button onclick="copyText('alt{i}')">Copy</button>
        </div>
        """
        )

    html_doc = f"""<!doctype html>
<html>
<head>
<meta charset="utf-8">
<title>{html.escape(title)}</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
  body {{ font-family: system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; margin: 24px; }}
  .toolbar {{ display:flex; gap:8px; align-items:center; margin-bottom:16px; flex-wrap: wrap; }}
  .item {{ border:1px solid #ddd; border-radius:12px; padding:12px; margin:12px 0; background:#fafafa; }}
  .item button {{ cursor:pointer; }}
  .alt {{ white-space: pre-wrap; word-wrap: break-word; margin:8px 0; background:#fff; padding:8px; border-radius:8px; border:1px solid #eee; }}
  .page {{ font-size: 12px; color:#666; }}
  .thumb {{ max-height:64px; max-width:96px; margin-right:8px; border-radius:8px; vertical-align: middle; }}
  .meta {{ display:flex; gap:8px; align-items:center; margin-bottom:6px; }}
  #toast {{ position:fixed; right:16px; bottom:16px; background:#222; color:#fff; padding:10px 14px; border-radius:10px; opacity:0; transition:opacity .2s; }}
  #toast.show {{ opacity:0.92; }}
</style>
</head>
<body>
  <h1>{html.escape(title)}</h1>
  <div class="toolbar">
    <button onclick="copyAll()">Copy all</button>
    <span id="count"></span>
  </div>
  {''.join(rows)}
  <div id="toast" role="status" aria-live="polite">Copied!</div>
<script>
  function showToast(msg) {{
    const t = document.getElementById('toast');
    t.textContent = msg || 'Copied!';
    t.classList.add('show');
    setTimeout(() => t.classList.remove('show'), 900);
  }}
  async function copyText(elemId) {{
    const el = document.getElementById(elemId);
    const txt = el ? el.textContent : '';
    try {{
      await navigator.clipboard.writeText(txt);
      showToast('Copied');
    }} catch (e) {{
      const r = document.createRange(); r.selectNode(el);
      const sel = window.getSelection(); sel.removeAllRanges(); sel.addRange(r);
      document.execCommand('copy'); sel.removeAllRanges();
      showToast('Copied');
    }}
  }}
  async function copyAll() {{
    const nodes = document.querySelectorAll('.alt');
    const block = Array.from(nodes).map(n => n.textContent.trim()).filter(Boolean).join('\\n\\n');
    try {{
      await navigator.clipboard.writeText(block);
      showToast('All copied');
    }} catch (e) {{
      const ta = document.createElement('textarea');
      ta.value = block; document.body.appendChild(ta); ta.select();
      document.execCommand('copy'); document.body.removeChild(ta);
      showToast('All copied');
    }}
  }}
  document.getElementById('count').textContent = document.querySelectorAll('.item').length + ' items';
</script>
</body>
</html>"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_doc)


def start_extraction_async(pdf_path: str, paths):
    """Start PDF image extraction in background."""
    import uuid
    logging.info(
        f"[sid={paths['base'].name}] worker START pdf_path={pdf_path}"
    )

    logging.info(f"[mem] start RSS={rss_mb():.0f}MB SOFT={SOFT_MEM_LIMIT_MB:.0f}MB")
    try:
        sid = paths["base"].name
        write_status(paths["img_status"], "Scanning PDF for images...")

        # Try PDFix first; on memory pressure/crash, fall back to PyMuPDF
        try:
            if not PDFIX_AVAILABLE:
                raise ImportError("PDFix SDK not available")
            images, page_map = extract_images_from_pdfix_per_session(
                pdf_path,
                str(paths["extracted"]),
                paths["img_progress"],
            )
        except (MemoryError, ImportError) as e:
            logging.exception(f"PDFix error ({type(e).__name__}); falling back to PyMuPDF")
            write_status(
                paths["img_status"], "PDFix unavailable; using PyMuPDF fallback…"
            )
            images, page_map = extract_images_with_pymupdf(
                pdf_path,
                str(paths["extracted"]),
                paths["img_progress"],
                sid=sid,
            )
        except Exception as e:
            msg = str(e).lower()
            if (
                ("out of memory" in msg)
                or ("cannot allocate" in msg)
                or ("segfault" in msg)
                or ("139" in msg)
                or ("pdfix" in msg and "not available" in msg)
            ):
                logging.exception("PDFix OOM/native crash; falling back to PyMuPDF")
                write_status(paths["img_status"], "PDFix failed; trying fallback…")
                images, page_map = extract_images_with_pymupdf(
                    pdf_path,
                    str(paths["extracted"]),
                    paths["img_progress"],
                    sid=sid,
                )
            else:
                # For any other error, still try PyMuPDF as fallback
                logging.exception(f"PDFix error: {e}; falling back to PyMuPDF")
                write_status(paths["img_status"], "PDFix error; trying fallback…")
                try:
                    images, page_map = extract_images_with_pymupdf(
                        pdf_path,
                        str(paths["extracted"]),
                        paths["img_progress"],
                        sid=sid,
                    )
                except Exception as fallback_error:
                    # If PyMuPDF also fails, raise the original error
                    raise e from fallback_error

        # Persist results for the frontend
        write_json(paths["page_map"], page_map)

        filenames = [img["filename"] for img in images]
        logging.info(f"[sid={sid}] extracted {len(filenames)} images")
        img_json = paths["base"] / "images.json"
        try:
            tmp = img_json.with_suffix(img_json.suffix + ".tmp")
            tmp.write_text(json.dumps(filenames), encoding="utf-8")
            os.replace(tmp, img_json)
        except Exception as e:
            logging.info(f"[sid={sid}] failed to write images.json: {e}")

        write_progress(paths["img_progress"], 100)  # ensure final tick
        write_status(paths["img_status"], "Done")  # <-- finish status
        logging.info(f"[sid={sid}] worker DONE images={len(images)}")
    except Exception as e:
        logging.exception("Extraction failed")
        (paths["base"] / "error.txt").write_text(f"{e}\n", encoding="utf-8")
        write_progress(paths["img_progress"], -1)
        write_status(paths["img_status"], "Error")


def background_generate_alt_text(images, lang, sid):
    """Generate alt text for images in background."""
    logging.info(
        f"Starting alt text generation for sid: {sid}, lang: {lang}, images: {len(images)}"
    )
    p = sess_paths(sid)
    p["output"].mkdir(parents=True, exist_ok=True)
    logging.info(f"Ensured output dir exists: {p['output']}")
    total = max(1, len(images))
    write_alt_progress_detail(p["alt_progress"], 0, 0, total)
    page_numbers = read_json(p["page_map"], {})
    doc = Document()
    entries = []

    def process_image(image_name):
        try:
            image_path = p["extracted"] / image_name
            page_number = page_numbers.get(image_name, "?")
            alt_text = (
                call_alttext_ai(str(image_path), lang) or "[No alt text generated]"
            )
            logging.info(f"Generated alt text for {image_name}: {alt_text[:50]}...")
            return {
                "alt": alt_text,
                "page": page_number,
                "thumb_rel": (
                    f"/pdf_processor/extracted_images/{sid}/{image_name}"
                    if image_path.exists()
                    else None
                ),
                "image_name": image_name,
            }
        except Exception as e:
            logging.info(f"Error processing {image_name}: {e}")
            return {
                "alt": "[Error generating alt text]",
                "page": page_number,
                "thumb_rel": None,
                "image_name": image_name,
            }

    with ThreadPoolExecutor(max_workers=MAX_ALT_TEXT_WORKERS) as alt_executor:
        results = []
        futures = [alt_executor.submit(process_image, img) for img in images]
        for i, future in enumerate(futures):
            entry = future.result()
            results.append(entry)
            done = i + 1
            percent = int((done / total) * 100)
            write_alt_progress_detail(p["alt_progress"], percent, done, total)
            logging.info(f"Progress: {done}/{total} images processed")

    for image_name, entry in zip(images, results):
        image_path = p["extracted"] / image_name
        if image_path.exists():
            try:
                with Image.open(str(image_path)) as img:
                    buf = BytesIO()
                    img.save(buf, format="PNG")
                    buf.seek(0)
                    doc.add_picture(buf, width=Inches(1))
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.add_run(f"Page {entry['page']}").italic = True
            except Exception as e:
                logging.info(f"⚠️ Could not add image: {e}")
                doc.add_paragraph("[Error displaying image]")
        else:
            doc.add_paragraph("[Image not found]")

        doc.add_paragraph(entry["alt"])
        entries.append(entry)

    docx_filename = f"alt_text_results_{sid}_{lang}.docx"
    word_doc_path = p["output"] / docx_filename
    doc.save(str(word_doc_path))
    logging.info(f"Saved Word file: {word_doc_path}")

    ready_path = p["output"] / f"ready_{sid}_{lang}.txt"
    ready_path.write_text("ready", encoding="utf-8")
    logging.info(f"Saved ready file: {ready_path}")

    copy_panel_path = p["output"] / f"copy_panel_{sid}.html"
    logging.info(f"Preparing to write copy panel: {copy_panel_path}")
    _write_copy_panel_html(
        str(copy_panel_path), entries, title=f"Alt Text Copy Panel ({lang.upper()})"
    )
    if copy_panel_path.exists():
        logging.info(f"Copy panel file exists after write: {copy_panel_path}")
    else:
        logging.info(f"Copy panel file NOT created: {copy_panel_path}")

